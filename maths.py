from random import randint
from datetime import datetime
import operator

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING, WD_BREAK

operators = [operator.add, operator.sub]
numbers = list(range(1, 21))
op_str = {
    operator.add: '+',
    operator.sub: '-'
}
quiz_per_line = 3


def select_operator():
    return operators[randint(0, 1)]


def select_number():
    n = randint(0, 9)
    if n < 8:
        return numbers[randint(0, 8)]
    return randint(9, numbers[-1])


def make_quiz(n):
    if n == 2:
        quiz = [select_number(), select_operator(), select_number()]
        val = eval_quiz(quiz)
        if val < 0:
            quiz, val = make_quiz(n)
    else:
        quiz, val = make_quiz(n - 1)
        op, number = select_operator(), select_number()
        while op(val, number) < 0:
            op, number = select_operator(), select_number()
        quiz = [number, op] + quiz
    return quiz, val


def eval_quiz(quiz):
    quiz1 = quiz[:]
    ret = quiz1.pop()
    while quiz1:
        op = quiz1.pop()
        ret = op(ret, quiz1.pop())
    return ret


def quiz_to_str(quiz):
    quiz1 = quiz[:]
    ret = str(quiz1.pop())
    while quiz1:
        ret += ' %s ' % op_str[quiz1.pop()]
        ret += str(quiz1.pop())
    return ret + ' ='


def make_sheet(num):
    quiz_list = []
    for i in range(num):
        num = randint(0, 6) < 5 and 2 or 3
        q, _ = make_quiz(num)
        quiz_list.append(q)
    return quiz_list


def format_sheet(sheet):
    i = 0
    ret = ''
    for quiz in sheet:
        str_quiz = quiz_to_str(quiz)
        ret += str_quiz
        ret += ' ' * (19 - len(str_quiz))
        i += 1
        if i == quiz_per_line:
            ret += '\n'
            i = 0
    return ret[-1] == '\n' and ret[:-1] or ret


def save_sheet_to_docx(sheet, doc, page_break_flag):
    p = doc.add_paragraph(sheet)
    p.line_spacing_rule = WD_LINE_SPACING.SINGLE
    # run.font.size = Pt(10)
    if page_break_flag:
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)


def init_doc():
    doc = Document()
    font = doc.styles['Normal'].font
    font.name = 'Source Code Pro'
    font.size = Pt(13)
    return doc


if __name__ == '__main__':
    import sys

    times = len(sys.argv) > 1 and int(sys.argv[1]) or 7
    num = len(sys.argv) > 2 and int(sys.argv[2]) or 100
    doc = init_doc()
    for i in range(times):
        sheet = format_sheet(make_sheet(num))
        save_sheet_to_docx(sheet, doc, i < times - 1)
    now = datetime.now()
    doc.save('quiz{}.docx'.format(now.strftime('%Y%m%d')))
